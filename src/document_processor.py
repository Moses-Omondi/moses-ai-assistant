#!/usr/bin/env python3
"""
Document Processor for DevSecOps AI Assistant
Processes various document types and converts them to searchable embeddings

Handles:
- PDF files (resumes, certifications)
- Word documents (project docs)
- Markdown files (README, documentation)
- Text files (notes, configurations)
- GitHub repositories (code, READMEs)
"""

import os
import logging
from pathlib import Path
from typing import List, Dict, Any
import chromadb
from chromadb.config import Settings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_community.embeddings import HuggingFaceEmbeddings
import PyPDF2
import docx
import markdown
from bs4 import BeautifulSoup
import git
import pandas as pd

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Processes documents and creates embeddings for RAG system
    Specialized for DevSecOps and AI Engineering content
    """
    
    def __init__(self, data_dir: str = None, db_path: str = None):
        if data_dir is None:
            data_dir = str(Path(__file__).parent.parent / "data")
        if db_path is None:
            db_path = str(Path(__file__).parent.parent / "data" / "chroma_db")
        self.data_dir = Path(data_dir)
        self.db_path = Path(db_path)
        
        # Initialize embeddings model (runs locally on M2)
        logger.info("Initializing embedding model...")
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'}  # M2 Mac optimization
        )
        
        # Initialize ChromaDB for vector storage
        self.client = chromadb.PersistentClient(path=str(self.db_path))
        
        # Create collection for Moses's professional knowledge
        self.collection_name = "moses_devsecops_knowledge"
        try:
            # Try to get existing collection
            self.collection = self.client.get_collection(self.collection_name)
            logger.info(f"Found existing collection: {self.collection_name}")
        except:
            # Create new collection if it doesn't exist
            self.collection = self.client.create_collection(
                name=self.collection_name,
                metadata={"description": "Moses Omondi's DevSecOps and AI Engineering expertise"}
            )
            logger.info(f"Created new collection: {self.collection_name}")
        
        # Text splitter for chunking documents
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,        # Good size for detailed technical content
            chunk_overlap=200,      # Maintain context between chunks
            separators=["\\n\\n", "\\n", " ", ""]
        )
    
    def process_pdf(self, file_path: Path) -> List[Document]:
        """Process PDF files (resume, certificates, documentation)"""
        logger.info(f"Processing PDF: {file_path}")
        documents = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text_content += page.extract_text() + "\\n"
                
                # Create document with metadata
                doc = Document(
                    page_content=text_content,
                    metadata={
                        "source": str(file_path),
                        "type": "pdf",
                        "category": self._categorize_document(file_path.name),
                        "pages": len(pdf_reader.pages)
                    }
                )
                documents.append(doc)
                logger.info(f"Extracted {len(pdf_reader.pages)} pages from {file_path.name}")
                
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
        
        return documents
    
    def process_docx(self, file_path: Path) -> List[Document]:
        """Process Word documents (project documentation)"""
        logger.info(f"Processing DOCX: {file_path}")
        documents = []
        
        try:
            doc = docx.Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\\n"
            
            document = Document(
                page_content=text_content,
                metadata={
                    "source": str(file_path),
                    "type": "docx",
                    "category": self._categorize_document(file_path.name),
                    "paragraphs": len(doc.paragraphs)
                }
            )
            documents.append(document)
            logger.info(f"Processed {len(doc.paragraphs)} paragraphs from {file_path.name}")
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
        
        return documents
    
    def process_markdown(self, file_path: Path) -> List[Document]:
        """Process Markdown files (READMEs, documentation)"""
        logger.info(f"Processing Markdown: {file_path}")
        documents = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
                
                # Convert markdown to HTML then to text for better processing
                html = markdown.markdown(md_content)
                soup = BeautifulSoup(html, 'html.parser')
                text_content = soup.get_text()
                
                document = Document(
                    page_content=text_content,
                    metadata={
                        "source": str(file_path),
                        "type": "markdown",
                        "category": self._categorize_document(file_path.name),
                        "original_format": "markdown"
                    }
                )
                documents.append(document)
                logger.info(f"Processed markdown file: {file_path.name}")
                
        except Exception as e:
            logger.error(f"Error processing Markdown {file_path}: {e}")
        
        return documents
    
    def process_text_file(self, file_path: Path) -> List[Document]:
        """Process plain text files"""
        logger.info(f"Processing text file: {file_path}")
        documents = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
                document = Document(
                    page_content=content,
                    metadata={
                        "source": str(file_path),
                        "type": "text",
                        "category": self._categorize_document(file_path.name)
                    }
                )
                documents.append(document)
                logger.info(f"Processed text file: {file_path.name}")
                
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
        
        return documents
    
    def process_github_repo(self, repo_url: str) -> List[Document]:
        """Clone and process GitHub repositories"""
        logger.info(f"Processing GitHub repo: {repo_url}")
        documents = []
        
        try:
            repo_name = repo_url.split('/')[-1].replace('.git', '')
            clone_path = self.data_dir / "github" / repo_name
            
            # Clone or update repository
            if clone_path.exists():
                logger.info(f"Updating existing repo: {repo_name}")
                repo = git.Repo(clone_path)
                repo.remotes.origin.pull()
            else:
                logger.info(f"Cloning repo: {repo_name}")
                git.Repo.clone_from(repo_url, clone_path)
            
            # Process README files and documentation
            for file_path in clone_path.rglob("*.md"):
                docs = self.process_markdown(file_path)
                for doc in docs:
                    doc.metadata["repo"] = repo_name
                    doc.metadata["repo_url"] = repo_url
                documents.extend(docs)
            
            # Process configuration files (YAML, JSON, etc.)
            config_extensions = ['.yml', '.yaml', '.json', '.toml', '.cfg']
            for ext in config_extensions:
                for file_path in clone_path.rglob(f"*{ext}"):
                    if file_path.stat().st_size < 50000:  # Skip very large files
                        docs = self.process_text_file(file_path)
                        for doc in docs:
                            doc.metadata["repo"] = repo_name
                            doc.metadata["repo_url"] = repo_url
                            doc.metadata["file_type"] = "configuration"
                        documents.extend(docs)
            
            logger.info(f"Processed {len(documents)} documents from {repo_name}")
            
        except Exception as e:
            logger.error(f"Error processing GitHub repo {repo_url}: {e}")
        
        return documents
    
    def _categorize_document(self, filename: str) -> str:
        """Categorize documents based on filename and content"""
        filename_lower = filename.lower()
        
        # DevSecOps and Security categories
        if any(keyword in filename_lower for keyword in ['security', 'sec', 'compliance', 'audit']):
            return "security"
        elif any(keyword in filename_lower for keyword in ['ci', 'cd', 'pipeline', 'deploy', 'jenkins', 'github-actions']):
            return "cicd"
        elif any(keyword in filename_lower for keyword in ['kubernetes', 'k8s', 'docker', 'container']):
            return "infrastructure"
        elif any(keyword in filename_lower for keyword in ['ml', 'ai', 'model', 'training', 'mlops']):
            return "ai_engineering"
        elif any(keyword in filename_lower for keyword in ['aws', 'cloud', 'cert', 'so3']):
            return "cloud_certification"
        elif any(keyword in filename_lower for keyword in ['resume', 'cv']):
            return "professional_profile"
        elif any(keyword in filename_lower for keyword in ['project', 'implementation']):
            return "project_documentation"
        else:
            return "general"
    
    def ingest_documents(self, documents: List[Document]):
        """Add documents to ChromaDB with embeddings"""
        if not documents:
            logger.warning("No documents to ingest")
            return
        
        logger.info(f"Ingesting {len(documents)} documents...")
        
        # Split documents into chunks
        chunks = []
        for doc in documents:
            doc_chunks = self.text_splitter.split_documents([doc])
            chunks.extend(doc_chunks)
        
        logger.info(f"Created {len(chunks)} chunks from {len(documents)} documents")
        
        # Prepare data for ChromaDB
        texts = [chunk.page_content for chunk in chunks]
        metadatas = [chunk.metadata for chunk in chunks]
        ids = [f"doc_{i}" for i in range(len(chunks))]
        
        # Generate embeddings and store in ChromaDB
        try:
            self.collection.add(
                documents=texts,
                metadatas=metadatas,
                ids=ids
            )
            logger.info(f"Successfully ingested {len(chunks)} chunks into ChromaDB")
        except Exception as e:
            logger.error(f"Error ingesting documents: {e}")
    
    def process_directory(self, directory_path: Path = None):
        """Process all documents in a directory"""
        if directory_path is None:
            directory_path = self.data_dir / "documents"
        
        logger.info(f"Processing directory: {directory_path}")
        all_documents = []
        
        if not directory_path.exists():
            logger.error(f"Directory does not exist: {directory_path}")
            return
        
        # Process different file types
        file_processors = {
            '.pdf': self.process_pdf,
            '.docx': self.process_docx,
            '.doc': self.process_docx,
            '.md': self.process_markdown,
            '.txt': self.process_text_file,
            '.yaml': self.process_text_file,
            '.yml': self.process_text_file,
            '.json': self.process_text_file
        }
        
        for file_path in directory_path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in file_processors:
                processor = file_processors[file_path.suffix.lower()]
                documents = processor(file_path)
                all_documents.extend(documents)
        
        # Ingest all documents
        if all_documents:
            self.ingest_documents(all_documents)
            logger.info(f"Processed {len(all_documents)} documents from {directory_path}")
        else:
            logger.warning(f"No supported documents found in {directory_path}")
    
    def get_collection_stats(self):
        """Get statistics about the knowledge base"""
        try:
            count = self.collection.count()
            logger.info(f"Knowledge base contains {count} document chunks")
            return {"total_chunks": count}
        except Exception as e:
            logger.error(f"Error getting collection stats: {e}")
            return {"total_chunks": 0}

if __name__ == "__main__":
    # Example usage
    processor = DocumentProcessor()
    
    # Process local documents
    processor.process_directory()
    
    # Show stats
    stats = processor.get_collection_stats()
    print(f"Knowledge base ready with {stats['total_chunks']} chunks!")
